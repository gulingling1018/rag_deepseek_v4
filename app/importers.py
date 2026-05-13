from collections import Counter, defaultdict
from dataclasses import dataclass, field
from functools import lru_cache
from html.parser import HTMLParser
from pathlib import Path
import re
import subprocess
import tempfile
import unicodedata
from urllib.parse import urlparse
from urllib.request import Request, urlopen

from charset_normalizer import from_bytes
from docx import Document as DocxDocument
from pypdf import PdfReader

from app.config import get_settings

TEXT_SUFFIXES = {".txt", ".text"}
MARKDOWN_SUFFIXES = {".md", ".markdown", ".mdx"}
WORD_SUFFIXES = {".docx"}
CODE_SUFFIXES = {
    ".c", ".h", ".cpp", ".hpp", ".cc", ".hh",
    ".py", ".js", ".ts", ".tsx", ".jsx",
    ".java", ".go", ".rs", ".swift", ".kt",
    ".json", ".yaml", ".yml", ".toml", ".ini",
    ".sh", ".bash", ".zsh", ".sql", ".html",
    ".css", ".scss", ".less", ".xml",
}

COMMON_CJK_HINTS = set(
    "的一是在不了有和人这中大为上个国我以要他时来用们生到作地于出就分对成会可主发年动同工也能下过子说产种面而方后多定行学法所民得经十三之进着等部度家电力里如水化高自二理起小物现实加量都两体制机当使点从业本去把性好应开它合还因由其些然前外天政四日那社义事平形相全表间样与关各线内数正心反你明看原或利比质气第向道命此变条只没结解问意建月公无系军很情者最立代想已通并提直题党程展五果料象员革位入常文总次品式活设及管特件长求老头基资边流路级少图山统接知较将组见计别她手角期根论运农指几九区强放决西被干做必战先回则任取据处理世蓝牙开发展编测试码主机控制器协议栈"
)


@dataclass
class ExtractedBlock:
    text: str
    block_type: str = "text"
    symbol_name: str | None = None
    section_path: list[str] = field(default_factory=list)
    page_number: int | None = None
    page_label: str | None = None
    line_start: int | None = None
    line_end: int | None = None
    paragraph_start: int | None = None
    paragraph_end: int | None = None
    source_uri: str | None = None
    location_label: str | None = None


@dataclass
class ExtractedDocument:
    text: str
    blocks: list[ExtractedBlock]
    source_format: str
    encoding: str | None = None
    page_count: int | None = None


@dataclass
class PdfPageData:
    page_number: int
    page_label: str
    text: str
    raw_lines: list[str]
    lines: list[str] = field(default_factory=list)
    used_ocr: bool = False


@dataclass
class PdfTocEntry:
    title: str
    target_page: int
    level: int
    path: list[str]
    entry_type: str = "section"


def normalize_text(text: str) -> str:
    normalized = unicodedata.normalize("NFKC", text)
    normalized = normalized.replace("\r\n", "\n").replace("\r", "\n")
    normalized = normalized.replace("\x00", "")
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    return normalized.strip()


def strip_sphinx_boilerplate(text: str) -> str:
    normalized = normalize_text(text)
    guard_markers = [
        "此文档对您有帮助吗",
        "Was this page helpful",
        "反馈已收到",
        "Thank you! We received your feedback",
        "利用Sphinx构建",
        "Built withSphinx",
        "PreviousNext",
        "上一页下一页",
    ]
    if not any(marker in normalized for marker in guard_markers):
        return normalized

    cut_markers = guard_markers + [
        "如果您有其他意见",
        "我们重视您的反馈",
        "Let us know how we can improve this page",
        "© 版权所有",
        "© Copyright",
        "Download HTML",
    ]
    cut = len(normalized)
    for marker in cut_markers:
        position = normalized.find(marker)
        if position >= 0:
            cut = min(cut, position)
    return normalize_text(normalized[:cut])


def slugify(value: str) -> str:
    slug = re.sub(r"[^\w\u4e00-\u9fff-]+", "-", value.strip().lower())
    slug = re.sub(r"-{2,}", "-", slug).strip("-")
    return slug or "section"


def clean_heading_text(text: str) -> str:
    cleaned = text.replace("\uf0c1", "").replace("", "")
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    return cleaned


def clean_pdf_markdown_line(line: str) -> str:
    cleaned = line.replace("_•_", "-").replace("•", "-")
    cleaned = clean_heading_text(cleaned)
    cleaned = re.sub(r"\s+GoBack$", "", cleaned)
    if cleaned == "GoBack":
        return ""
    return cleaned.strip()


def is_pdf_markdown_noise_line(line: str) -> bool:
    stripped = clean_pdf_markdown_line(line)
    if not stripped:
        return False
    if is_pdf_footer_or_header_line(stripped):
        return True
    if re.fullmatch(r"Espressif Systems\s+\d+", stripped):
        return True
    if re.fullmatch(r"ESP32(?:-[A-Z0-9]+)?\s+(?:TRM|Datasheet)\s+\(Version\s+[^)]+\)", stripped, flags=re.IGNORECASE):
        return True
    if "Submit Documentation Feedback" in stripped:
        return True
    if "documentation_feedback" in stripped:
        return True
    if re.fullmatch(r"Chapter\s+\d+\s+.+", stripped, flags=re.IGNORECASE):
        return True
    return False


def is_pdf_footer_or_header_line(line: str) -> bool:
    stripped = clean_heading_text(line)
    if not stripped:
        return True
    footer_patterns = [
        r"乐鑫信息科技\s+\d+\s+ESP32(?:-[A-Z0-9]+)?(?:\s*系列芯片)?(?:技术规格书|技术参考手册|硬件设计指南)?\s*v?[\d.]*",
        r"ESP32(?:-[A-Z0-9]+)?(?:\s*系列芯片)?(?:技术规格书|技术参考手册|硬件设计指南)?\s*v?[\d.]*\s+乐鑫信息科技\s+\d+",
        r"Espressif Systems\s+\d+",
        r"ESP32(?:-[A-Z0-9]+)?\s+(?:TRM|Datasheet)\s+\(Version\s+[^)]+\)",
    ]
    return any(re.fullmatch(pattern, stripped, flags=re.IGNORECASE) for pattern in footer_patterns)


def strip_pdf_noise_from_text(text: str) -> str:
    lines: list[str] = []
    for raw_line in text.splitlines():
        line = clean_pdf_markdown_line(raw_line)
        if not line or is_pdf_markdown_noise_line(line):
            continue
        line = re.sub(
            r"\s*乐鑫信息科技\s+\d+\s+ESP32(?:-[A-Z0-9]+)?(?:\s*系列芯片)?(?:技术规格书|技术参考手册|硬件设计指南)?\s*v?[\d.]*\s*",
            " ",
            line,
            flags=re.IGNORECASE,
        )
        line = re.sub(
            r"\s*[iI]?\s*乐鑫信息科技\s+\d+\s*[iI]?\s+ESP32(?:-[A-Z0-9]+)?(?:\s*系列芯片)?(?:技术规格书|技术参考手册|硬件设计指南)?\s*v?[\d.]*\s*",
            " ",
            line,
            flags=re.IGNORECASE,
        )
        line = re.sub(
            r"\s*ESP32(?:-[A-Z0-9]+)?(?:\s*系列芯片)?(?:技术规格书|技术参考手册|硬件设计指南)?\s*v?[\d.]*\s+乐鑫信息科技\s+\d+\s*",
            " ",
            line,
            flags=re.IGNORECASE,
        )
        line = clean_heading_text(line)
        if not line or is_pdf_markdown_noise_line(line):
            continue
        lines.append(line)
    return "\n".join(lines).strip()


def is_repeated_markdown_table_row(line: str) -> bool:
    stripped = line.strip()
    if not stripped.startswith("|") or not stripped.endswith("|"):
        return False
    if re.fullmatch(r"\|(?:\s*:?-{3,}:?\s*\|)+", stripped):
        return False
    cells = [re.sub(r"<br\s*/?>", " ", cell, flags=re.IGNORECASE).strip(" *_") for cell in stripped.strip("|").split("|")]
    cells = [cell for cell in cells if cell]
    if len(cells) < 2:
        return False
    return len(set(cells)) == 1


def clean_pdf_table_text(text: str) -> str:
    lines = []
    for line in strip_pdf_noise_from_text(text).splitlines():
        if is_repeated_markdown_table_row(line):
            continue
        lines.append(line)
    if lines and lines[0] == "Table" and any(line.startswith("|") for line in lines[1:]):
        lines = lines[1:]
    return "\n".join(lines).strip()


def reduce_figure_ocr_text(text: str) -> str:
    single_line = clean_heading_text(text.replace("\n", " "))
    figure_match = re.search(
        r"((?:Figure|Fig\.)\s+\d+(?:[\.-]\d+)*\.?\s+.+|(?:\*\*)?图(?:\*\*)?\s*\d+[-.]\d+\.?\s+.+)$",
        single_line,
        flags=re.IGNORECASE,
    )
    if not figure_match:
        return text

    prefix = single_line[: figure_match.start()].strip()
    if len(prefix) < 80:
        return text

    tokens = prefix.split()
    if not tokens:
        return text

    digit_tokens = sum(1 for token in tokens if re.fullmatch(r"\d+", token))
    pin_like_tokens = sum(1 for token in tokens if re.fullmatch(r"[A-Z0-9_./+-]{2,}", token))
    sentence_marks = len(re.findall(r"[.。？！!?;；:：]", prefix))
    if len(tokens) >= 18 and (digit_tokens + pin_like_tokens) / len(tokens) > 0.45 and sentence_marks <= 2:
        return figure_match.group(1).strip()
    return text


def is_low_value_pdf_text(text: str) -> bool:
    stripped = clean_heading_text(re.sub(r"[*_`#]+", "", text))
    if not stripped:
        return True
    if stripped in {"目录", "表格", "插图", "Table", "Contents", "List of Tables", "List of Figures"}:
        return True
    if "免责声明和版权公告" in stripped:
        return True
    if len(stripped) < 80:
        return True
    return False


def pdf_text_ends_mid_sentence(text: str) -> bool:
    stripped = text.strip()
    if not stripped:
        return False
    if re.search(r"[.!?。？！][\"'”’)\]]?$", stripped):
        return False
    if re.search(r"\b(?:the|a|an|of|to|with|from|for|and|or|on|in|by|as|is|are|USB)$", stripped, flags=re.IGNORECASE):
        return True
    return len(stripped) < 180


def pdf_text_starts_continuation(text: str) -> bool:
    stripped = text.strip()
    if not stripped or stripped.startswith(("-", "•", "|")):
        return False
    return bool(re.match(r"^(?:[a-z]|and\b|or\b|which\b|that\b|with\b|from\b|to\b)", stripped, flags=re.IGNORECASE))


def should_merge_pdf_text_blocks(previous: ExtractedBlock, current: ExtractedBlock) -> bool:
    if previous.block_type != "text" or current.block_type != "text":
        return False
    if previous.section_path != current.section_path:
        return False
    if previous.page_number is not None and current.page_number is not None:
        if current.page_number - previous.page_number not in (0, 1):
            return False
    if len(previous.text) + len(current.text) > 2200:
        return False
    return pdf_text_ends_mid_sentence(previous.text) or pdf_text_starts_continuation(current.text)


def is_probable_markdown_toc_page(page_number: int, lines: list[str]) -> bool:
    meaningful = [clean_pdf_markdown_line(line) for line in lines if clean_pdf_markdown_line(line)]
    if not meaningful or page_number > 20:
        return False
    normalized_first = re.sub(r"[*_#]+", "", meaningful[0]).strip()
    if normalized_first in {"目录", "表格", "插图", "Contents", "List of Tables", "List of Figures"}:
        return True
    toc_like = sum(1 for line in meaningful if re.search(r"\s+\d{1,4}$", line) and len(line) < 120)
    return toc_like >= 6


def sanitize_pdf_section_path(section_path: list[str]) -> list[str]:
    cleaned_path: list[str] = []
    for item in section_path:
        cleaned = strip_pdf_noise_from_text(item)
        cleaned = re.sub(r"[*_`#]+", "", cleaned).strip()
        if not cleaned or cleaned in {"目录", "表格", "插图", "Table", "Contents", "List of Tables", "List of Figures"}:
            continue
        if re.fullmatch(r"\d+\s+\d+", cleaned) or re.match(r"^\d+\s*:", cleaned):
            continue
        if len(cleaned) > 100 and not extract_numeric_heading_key(cleaned):
            continue
        if cleaned not in cleaned_path:
            cleaned_path.append(cleaned)
    return cleaned_path


def postprocess_pdf_blocks(blocks: list[ExtractedBlock]) -> list[ExtractedBlock]:
    processed: list[ExtractedBlock] = []
    pending_short: ExtractedBlock | None = None

    def append_block(block: ExtractedBlock) -> None:
        nonlocal pending_short
        if block.block_type == "table":
            if pending_short:
                processed.append(pending_short)
                pending_short = None
            processed.append(block)
            return

        if len(block.text) < 80:
            if pending_short and pending_short.page_number == block.page_number and pending_short.section_path == block.section_path:
                pending_short.text = f"{pending_short.text} {block.text}".strip()
                if len(pending_short.text) >= 80:
                    processed.append(pending_short)
                    pending_short = None
            elif pending_short:
                processed.append(pending_short)
                pending_short = block
            else:
                pending_short = block
            return

        if pending_short and pending_short.page_number == block.page_number and pending_short.section_path == block.section_path:
            block.text = f"{pending_short.text} {block.text}".strip()
            pending_short = None
        elif pending_short:
            processed.append(pending_short)
            pending_short = None
        if processed and should_merge_pdf_text_blocks(processed[-1], block):
            previous = processed[-1]
            previous.text = f"{previous.text} {block.text}".strip()
            if previous.page_label and block.page_label and previous.page_label != block.page_label:
                previous.location_label = f"{previous.page_label} - {block.page_label}"
            return
        processed.append(block)

    for block in blocks:
        text = clean_pdf_table_text(block.text) if block.block_type == "table" else strip_pdf_noise_from_text(block.text)
        if block.block_type != "table":
            text = reduce_figure_ocr_text(text)
        if block.block_type == "table":
            if not text or "|" not in text:
                continue
        elif is_low_value_pdf_text(text):
            continue

        block.text = text
        block.section_path = sanitize_pdf_section_path(block.section_path)
        append_block(block)

    if pending_short and not is_low_value_pdf_text(pending_short.text):
        processed.append(pending_short)
    return processed


def text_readability_score(text: str) -> float:
    chars = [char for char in text if not char.isspace()]
    if not chars:
        return 0.0

    good = 0
    common_punct = set(".,;:!?()[]{}<>-_+/\\|@#$%^&*~`'\"，。；：？！、（）《》【】“”‘’·…")
    for char in chars:
        if char in common_punct:
            good += 1
            continue
        if char.isascii() and (char.isalnum() or char in common_punct):
            good += 1
            continue
        if "\u4e00" <= char <= "\u9fff":
            good += 1
            continue

    return good / len(chars)


@lru_cache(maxsize=1)
def get_rapidocr_engine():
    from rapidocr_onnxruntime import RapidOCR

    return RapidOCR()


def extract_text_with_ocr(image_path: Path) -> str:
    lines = extract_ocr_lines(image_path)
    return normalize_text("\n".join(lines))


def extract_ocr_lines(image_path: Path) -> list[str]:
    engine = get_rapidocr_engine()
    result, _ = engine(str(image_path))
    if not result:
        return []

    items = []
    for item in result:
        points = item[0]
        text = clean_heading_text(str(item[1]))
        if not text:
            continue
        xs = [point[0] for point in points]
        ys = [point[1] for point in points]
        items.append(
            {
                "text": text,
                "x": min(xs),
                "x2": max(xs),
                "y": sum(ys) / len(ys),
                "h": max(ys) - min(ys),
            }
        )

    items.sort(key=lambda entry: (entry["y"], entry["x"]))
    if not items:
        return []

    rows: list[list[dict]] = []
    current_row = [items[0]]
    current_center = items[0]["y"]
    current_height = max(items[0]["h"], 1.0)

    for item in items[1:]:
        threshold = max(current_height * 0.6, 12.0)
        if abs(item["y"] - current_center) <= threshold:
            current_row.append(item)
            current_center = (current_center * (len(current_row) - 1) + item["y"]) / len(current_row)
            current_height = max(current_height, item["h"])
        else:
            rows.append(current_row)
            current_row = [item]
            current_center = item["y"]
            current_height = max(item["h"], 1.0)
    rows.append(current_row)

    lines: list[str] = []
    for row in rows:
        row.sort(key=lambda entry: entry["x"])
        parts: list[str] = []
        previous_x2 = None
        for entry in row:
            if previous_x2 is not None and entry["x"] - previous_x2 > 36:
                parts.append("  ")
            parts.append(entry["text"])
            previous_x2 = entry["x2"]
        line = clean_heading_text("".join(parts))
        if line:
            lines.append(line)
    return lines


def render_pdf_page_image(pdf_path: Path, page_number: int, dpi: int) -> Path | None:
    with tempfile.TemporaryDirectory(prefix="rag_pdf_ocr_") as tmpdir:
        prefix = Path(tmpdir) / "page"
        command = [
            "pdftoppm",
            "-f",
            str(page_number),
            "-l",
            str(page_number),
            "-r",
            str(dpi),
            "-png",
            str(pdf_path),
            str(prefix),
        ]
        subprocess.run(command, check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        images = sorted(Path(tmpdir).glob("page-*.png"))
        if not images:
            return None

        persistent = Path(tempfile.gettempdir()) / f"rag_pdf_ocr_page_{page_number}.png"
        persistent.write_bytes(images[0].read_bytes())
        return persistent


def select_pdf_page_text(pdf_path: Path, page_number: int, native_text: str) -> tuple[str, list[str], bool]:
    settings = get_settings()
    native_lines = [clean_heading_text(line) for line in native_text.splitlines() if clean_heading_text(line)]
    if not settings.rag_pdf_ocr_enabled:
        return native_text, native_lines, False

    mode = settings.rag_pdf_ocr_mode.lower()
    native_score = text_readability_score(native_text)
    should_try_ocr = mode == "always"
    if mode == "auto":
        should_try_ocr = (
            not native_text.strip()
            or native_score < settings.rag_pdf_ocr_quality_threshold
        )
    if not should_try_ocr:
        return native_text, native_lines, False

    image_path = None
    try:
        image_path = render_pdf_page_image(pdf_path, page_number, settings.rag_pdf_ocr_dpi)
        if not image_path:
            return native_text, native_lines, False
        ocr_lines = extract_ocr_lines(image_path)
        ocr_text = normalize_text("\n".join(ocr_lines))
    except Exception:
        return native_text, native_lines, False
    finally:
        if image_path and image_path.exists():
            image_path.unlink(missing_ok=True)

    if not ocr_text.strip():
        return native_text, native_lines, False

    ocr_score = text_readability_score(ocr_text)
    if mode == "always":
        return ocr_text, ocr_lines, True

    if ocr_score > native_score + 0.08:
        return ocr_text, ocr_lines, True
    if ocr_score >= native_score and len(ocr_text) > max(len(native_text), 1) * 1.2:
        return ocr_text, ocr_lines, True

    return native_text, native_lines, False


def decode_text_bytes(raw: bytes) -> tuple[str, str | None]:
    candidates: list[tuple[float, str, str]] = []

    def score_text(text: str) -> float:
        readability = text_readability_score(text)
        chars = [char for char in text if not char.isspace()]
        if not chars:
            return readability
        cjk_chars = [char for char in chars if "\u4e00" <= char <= "\u9fff"]
        common_hits = sum(1 for char in cjk_chars if char in COMMON_CJK_HINTS)
        common_ratio = common_hits / len(cjk_chars) if cjk_chars else 0.0
        ascii_ratio = sum(1 for char in chars if char.isascii()) / len(chars)
        return readability + common_ratio * 0.35 + ascii_ratio * 0.05

    detected = from_bytes(raw).best()
    if detected and detected.encoding:
        text = normalize_text(str(detected))
        candidates.append((score_text(text), text, detected.encoding))

    for encoding in ("utf-8", "utf-8-sig", "gb18030", "gbk", "big5", "latin-1"):
        try:
            text = normalize_text(raw.decode(encoding))
        except UnicodeDecodeError:
            continue
        candidates.append((score_text(text), text, encoding))

    if candidates:
        candidates.sort(key=lambda item: item[0], reverse=True)
        _, text, encoding = candidates[0]
        return text, encoding

    return normalize_text(raw.decode("utf-8", errors="replace")), "utf-8-replace"


def build_paragraph_blocks(
    paragraphs: list[tuple[str, dict]],
    *,
    max_chars: int = 1200,
    max_items: int = 12,
) -> list[ExtractedBlock]:
    blocks: list[ExtractedBlock] = []
    current_texts: list[str] = []
    current_meta: dict | None = None

    def flush():
        nonlocal current_texts, current_meta
        if not current_texts or current_meta is None:
            return
        text = "\n\n".join(item for item in current_texts if item.strip()).strip()
        if text:
            location_label = current_meta.get("location_label")
            if current_meta.get("paragraph_start") is not None and current_meta.get("paragraph_end") is not None:
                start = current_meta["paragraph_start"]
                end = current_meta["paragraph_end"]
                location_label = f"段落 {start}" if start == end else f"段落 {start}-{end}"
            elif current_meta.get("line_start") is not None and current_meta.get("line_end") is not None:
                start = current_meta["line_start"]
                end = current_meta["line_end"]
                location_label = f"行 {start}" if start == end else f"行 {start}-{end}"
            blocks.append(
                ExtractedBlock(
                    text=text,
                    block_type=current_meta.get("block_type", "text"),
                    symbol_name=current_meta.get("symbol_name"),
                    section_path=list(current_meta.get("section_path", [])),
                    page_number=current_meta.get("page_number"),
                    page_label=current_meta.get("page_label"),
                    line_start=current_meta.get("line_start"),
                    line_end=current_meta.get("line_end"),
                    paragraph_start=current_meta.get("paragraph_start"),
                    paragraph_end=current_meta.get("paragraph_end"),
                    source_uri=current_meta.get("source_uri"),
                    location_label=location_label,
                )
            )
        current_texts = []
        current_meta = None

    for text, meta in paragraphs:
        cleaned = normalize_text(text)
        if not cleaned:
            continue

        if current_meta is None:
            current_meta = dict(meta)
            current_texts = [cleaned]
            continue

        same_scope = (
            current_meta.get("page_number") == meta.get("page_number")
            and current_meta.get("section_path") == meta.get("section_path")
            and current_meta.get("source_uri") == meta.get("source_uri")
        )
        candidate = "\n\n".join(current_texts + [cleaned])
        if same_scope and len(candidate) <= max_chars and len(current_texts) < max_items:
            current_texts.append(cleaned)
            if meta.get("line_end") is not None:
                current_meta["line_end"] = meta.get("line_end")
            if meta.get("paragraph_end") is not None:
                current_meta["paragraph_end"] = meta.get("paragraph_end")
            if meta.get("location_label"):
                current_meta["location_label"] = meta["location_label"]
            continue

        flush()
        current_meta = dict(meta)
        current_texts = [cleaned]

    flush()
    return blocks


def build_line_blocks(
    lines: list[str],
    *,
    filename: str,
    source_uri: str | None = None,
    block_type: str = "text",
    max_chars: int = 1400,
    max_lines: int = 80,
) -> list[ExtractedBlock]:
    blocks: list[ExtractedBlock] = []
    current_section: list[str] = []
    pending: list[tuple[int, str]] = []
    current_symbol: str | None = None
    in_fenced_code = False
    code_fence = ""
    code_start_line: int | None = None
    code_pending: list[tuple[int, str]] = []

    def detect_markdown_heading(line: str) -> tuple[int, str] | None:
        match = re.match(r"^(#{1,6})\s+(.+)$", line.strip())
        if not match:
            return None
        return len(match.group(1)), clean_heading_text(match.group(2))

    def detect_code_symbol(line: str) -> str | None:
        patterns = [
            r"^\s*(class|struct|enum|interface)\s+([A-Za-z_][\w]*)",
            r"^\s*def\s+([A-Za-z_][\w]*)\s*\(",
            r"^\s*function\s+([A-Za-z_][\w]*)\s*\(",
            r"^\s*func\s+([A-Za-z_][\w]*)\s*\(",
            r"^\s*fn\s+([A-Za-z_][\w]*)\s*\(",
            r"^\s*(?:[A-Za-z_][\w\s\*]+)\s+([A-Za-z_][\w]*)\s*\([^;]*\)\s*\{?$",
        ]
        stripped = line.strip()
        if not stripped or stripped.startswith(("//", "#", "*", "/*")):
            return None
        for pattern in patterns:
            match = re.match(pattern, stripped)
            if match:
                return match.groups()[-1]
        return None

    def flush():
        nonlocal pending
        if not pending:
            return
        start = pending[0][0]
        end = pending[-1][0]
        text = "\n".join(item[1] for item in pending).strip()
        if text:
            section = list(current_section)
            if current_symbol and (not section or section[-1] != current_symbol):
                section = section + [current_symbol]
            label = f"{filename}:{start}-{end}"
            blocks.append(
                ExtractedBlock(
                    text=text,
                    block_type=block_type,
                    symbol_name=current_symbol,
                    section_path=section,
                    line_start=start,
                    line_end=end,
                    source_uri=source_uri,
                    location_label=label,
                )
            )
        pending = []

    def flush_code(end_line: int | None = None):
        nonlocal code_pending, code_start_line
        if not code_pending:
            code_start_line = None
            return
        start = code_start_line or code_pending[0][0]
        end = end_line or code_pending[-1][0]
        text = "\n".join(item[1] for item in code_pending).strip("\n")
        if text.strip():
            blocks.append(
                ExtractedBlock(
                    text=text,
                    block_type="code",
                    section_path=list(current_section),
                    line_start=start,
                    line_end=end,
                    source_uri=source_uri,
                    location_label=f"{filename}:{start}-{end}",
                )
            )
        code_pending = []
        code_start_line = None

    for index, raw_line in enumerate(lines, start=1):
        line = raw_line.rstrip("\n")
        fence_match = re.match(r"^\s*(```+|~~~+)", line)
        if fence_match:
            fence = fence_match.group(1)
            if in_fenced_code and fence.startswith(code_fence[:3]):
                flush_code(index)
                in_fenced_code = False
                code_fence = ""
                continue
            if not in_fenced_code:
                flush()
                in_fenced_code = True
                code_fence = fence
                code_start_line = index + 1
                code_pending = []
                continue

        if in_fenced_code:
            code_pending.append((index, line))
            if sum(len(item[1]) + 1 for item in code_pending) >= max_chars:
                flush_code(index)
                code_start_line = index + 1
            continue

        heading = detect_markdown_heading(line)
        if heading:
            flush()
            level, title = heading
            current_section = current_section[: level - 1]
            current_section.append(title)
            current_symbol = None
            continue

        symbol = detect_code_symbol(line) if block_type == "code" else None
        if symbol:
            flush()
            current_symbol = symbol

        pending.append((index, line))
        char_count = sum(len(item[1]) + 1 for item in pending)
        if len(pending) >= max_lines or char_count >= max_chars:
            flush()

    if in_fenced_code:
        flush_code(lines[-1][0] if lines and isinstance(lines[-1], tuple) else len(lines))
    flush()
    return blocks


def detect_pdf_heading(line: str) -> str | None:
    stripped = clean_heading_text(line.strip())
    if not stripped:
        return None
    if is_pdf_footer_or_header_line(stripped):
        return None
    if re.fullmatch(r"\d+\s+\d+", stripped) or re.match(r"^\d+\s*:", stripped):
        return None
    if re.match(r"^\d+\s+(?:Mbit/s|Kbit/s|Gbit/s|Hz|kHz|MHz|GHz|ns|us|µs|ms|s|V|mV|A|mA|dB|bytes?)\b", stripped, flags=re.IGNORECASE):
        return None
    if re.match(r"^\d+\s+\S+", stripped) and len(stripped) > 90:
        return None
    if re.match(r"^\d+\s+\S+", stripped) and len(stripped) > 55 and re.search(r"[，。,;；]", stripped):
        return None
    patterns = [
        r"^\d+(?:\.\d+)+\s*\S+",
        r"^\d+\s+\S+",
        r"^\d+[\u4e00-\u9fff]\S*",
        r"^第[一二三四五六七八九十百零0-9]+[章节部分篇]\s*\S*",
        r"^(表|图)\s*\d+[-.]\d+\s*\S+",
        r"^(chapter|section)\s+\d+[:.\s-]+\S+",
    ]
    for pattern in patterns:
        if re.match(pattern, stripped, flags=re.IGNORECASE):
            return stripped
    return None


def heading_level(heading: str) -> int:
    numbered = re.match(r"^(\d+(?:\.\d+)*)", heading)
    if numbered:
        return min(len(numbered.group(1).split(".")), 6)
    appendix = re.match(r"^(?:附录\s*)?([A-Z])(?:\.(\d+(?:\.\d+)*)?)?", heading, flags=re.IGNORECASE)
    if appendix:
        appendix_parts = appendix.group(2)
        return 1 if not appendix_parts else min(1 + len(appendix_parts.split(".")), 6)
    if heading.startswith(("表", "图")):
        return 6
    return 1


def update_section_path(current: list[str], heading: str) -> list[str]:
    title = clean_heading_text(heading)
    level = heading_level(title)
    base = current[: max(level - 1, 0)]
    base.append(title)
    return base


def is_pdf_footer_line(line: str, page_number: int) -> bool:
    stripped = clean_heading_text(line)
    if not stripped:
        return True
    if "反馈文档意见" in stripped or "乐鑫信息科技" in stripped:
        return True
    if "技术规格书" in stripped and "ESP32" in stripped:
        return True
    if stripped == str(page_number):
        return True
    if re.fullmatch(r"\d{1,3}", stripped):
        return True
    if re.fullmatch(r"第\s*\d+\s*页", stripped):
        return True
    return False


def is_pdf_toc_entry(line: str) -> bool:
    stripped = clean_heading_text(line)
    if not stripped:
        return False
    if re.search(r"\b\d{1,3}$", stripped) and (
        re.search(r"\d+(?:\.\d+)*", stripped) or re.match(r"^(附录\s*[A-Z]|[A-Z]\.)", stripped, flags=re.IGNORECASE)
    ):
        return True
    if stripped.startswith(("插图", "目录", "表格")) and re.search(r"\b\d{1,3}$", stripped):
        return True
    return False


def is_pdf_toc_page(page_number: int, lines: list[str]) -> bool:
    meaningful = [clean_heading_text(line) for line in lines if clean_heading_text(line)]
    if not meaningful:
        return False
    if any(any(marker in line for marker in ("目录", "插图", "表格")) for line in meaningful[:6]):
        return True
    if page_number > 15:
        return False
    toc_like = sum(1 for line in meaningful if is_pdf_toc_entry(line))
    return toc_like >= 4


def parse_pdf_toc_page(page_number: int, lines: list[str]) -> list[ExtractedBlock]:
    entries = [clean_heading_text(line) for line in lines if not is_pdf_footer_line(line, page_number)]
    toc_lines = [line for line in entries if line and line != "目录"]
    chunks: list[ExtractedBlock] = []
    for start in range(0, len(toc_lines), 6):
        text = "\n".join(toc_lines[start : start + 6]).strip()
        if text:
            chunks.append(
                ExtractedBlock(
                    text=text,
                    block_type="toc",
                    section_path=["目录"],
                    page_number=page_number,
                    page_label=f"第 {page_number} 页",
                    location_label=f"第 {page_number} 页",
                )
            )
    return chunks


def is_pdf_table_title(line: str) -> bool:
    stripped = clean_heading_text(line)
    return re.search(r"(表|table)\s*\d+[-.]\d+\s*\S+", stripped, flags=re.IGNORECASE) is not None


def extract_inline_table_title(line: str) -> tuple[str | None, str | None]:
    stripped = clean_heading_text(line)
    match = re.search(r"(表|table)\s*\d+[-.]\d+\s*\S.*", stripped, flags=re.IGNORECASE)
    if not match:
        return None, None
    prefix = stripped[: match.start()].strip()
    table_title = stripped[match.start() :].strip()
    return prefix or None, table_title or None


def looks_like_table_row(line: str) -> bool:
    stripped = clean_heading_text(line)
    if not stripped:
        return False
    if is_pdf_table_title(stripped):
        return True
    if detect_pdf_heading(stripped):
        return False
    if "  " in line:
        return True
    if len(re.findall(r"[-+]?\d+(?:\.\d+)?", stripped)) >= 2:
        return True
    headers = ["参数", "说明", "最小值", "最大值", "典型值", "单位"]
    if sum(1 for header in headers if header in stripped) >= 2:
        return True
    return False


def looks_like_table_header_row(line: str) -> bool:
    stripped = clean_heading_text(line)
    if not stripped or is_pdf_table_title(stripped):
        return False
    headers = [
        "参数",
        "说明",
        "最小值",
        "最大值",
        "典型值",
        "单位",
        "测试条件",
        "测试标准",
        "工作模式",
    ]
    if sum(1 for header in headers if header in stripped) >= 2:
        return True
    if len(re.findall(r"[-+]?\d+(?:\.\d+)?", stripped)) <= 1 and len(stripped) <= 80 and " " in stripped:
        return True
    return False


def normalize_pdf_noise_signature(line: str) -> str:
    stripped = clean_heading_text(line)
    stripped = re.sub(r"第\s*\d+\s*页", " ", stripped)
    stripped = re.sub(r"\b\d{1,3}\b", "#", stripped)
    stripped = re.sub(r"\s+", "", stripped)
    return stripped


def collect_pdf_margin_noise(pages: list[PdfPageData]) -> tuple[set[str], set[str]]:
    header_counts: Counter[str] = Counter()
    footer_counts: Counter[str] = Counter()
    min_hits = 3 if len(pages) < 20 else max(3, len(pages) // 10)

    for page in pages:
        meaningful = [clean_heading_text(line) for line in page.raw_lines if clean_heading_text(line)]
        if not meaningful:
            continue
        for line in meaningful[:2]:
            signature = normalize_pdf_noise_signature(line)
            if len(signature) >= 4 and not any(marker in line for marker in ("目录", "表格", "插图")):
                header_counts[signature] += 1
        for line in meaningful[-3:]:
            signature = normalize_pdf_noise_signature(line)
            if len(signature) >= 4 and not any(marker in line for marker in ("目录", "表格", "插图")):
                footer_counts[signature] += 1

    header_noise = {signature for signature, count in header_counts.items() if count >= min_hits}
    footer_noise = {signature for signature, count in footer_counts.items() if count >= min_hits}
    return header_noise, footer_noise


def strip_pdf_page_noise(
    page: PdfPageData,
    *,
    header_noise: set[str],
    footer_noise: set[str],
) -> list[str]:
    meaningful = [clean_heading_text(line) for line in page.raw_lines if clean_heading_text(line)]
    total = len(meaningful)
    filtered: list[str] = []
    for index, line in enumerate(meaningful):
        signature = normalize_pdf_noise_signature(line)
        is_margin_noise = (
            (index < 2 and signature in header_noise)
            or (index >= max(total - 3, 0) and signature in footer_noise)
        )
        if is_margin_noise or is_pdf_footer_line(line, page.page_number):
            continue
        filtered.append(line)
    return filtered


def normalize_pdf_toc_title(title: str, entry_type: str) -> str:
    normalized = clean_heading_text(title).strip(" .·•")
    normalized = re.sub(r"\s+", " ", normalized)
    if entry_type == "table":
        normalized = re.sub(r"^表?\s*(\d+[-.]\d+)\s*", r"表\1. ", normalized)
    elif entry_type == "figure":
        normalized = re.sub(r"^图?\s*(\d+[-.]\d+)\s*", r"图\1. ", normalized)
    return normalized.strip()


def parse_pdf_toc_line(line: str, entry_type: str) -> tuple[str, int] | None:
    stripped = clean_heading_text(line)
    match = re.match(r"^(?P<title>.+?)\s+(\d{1,3})$", stripped)
    if not match:
        return None
    title = normalize_pdf_toc_title(match.group("title"), entry_type)
    target_page = int(match.group(2))
    if not title or target_page <= 0:
        return None
    return title, target_page


def is_probable_toc_fragment(line: str) -> bool:
    stripped = clean_heading_text(line)
    if not stripped:
        return False
    if stripped in {"目录", "表格", "插图"}:
        return False
    if re.search(r"\b\d{1,3}$", stripped):
        return False
    if re.match(r"^(\d+(?:\.\d+)*|附录\s*[A-Z]|[A-Z]\.)", stripped, flags=re.IGNORECASE):
        return True
    return stripped.startswith(("封装", "相关文档和资源", "修订历史"))


def toc_entry_level(title: str) -> int:
    numbered = re.match(r"^(\d+(?:\.\d+)*)", title)
    if numbered:
        return min(len(numbered.group(1).split(".")), 6)
    appendix = re.match(r"^(?:附录\s*)?([A-Z])(?:\.(\d+(?:\.\d+)*)?)?", title, flags=re.IGNORECASE)
    if appendix:
        appendix_parts = appendix.group(2)
        return 1 if not appendix_parts else min(1 + len(appendix_parts.split(".")), 6)
    return 1


def build_pdf_toc_entries(pages: list[PdfPageData]) -> list[PdfTocEntry]:
    entries: list[PdfTocEntry] = []
    section_stack: list[str] = []
    current_mode = "section"

    def append_entry(title: str, target_page: int) -> None:
        nonlocal section_stack
        if current_mode == "section":
            level = toc_entry_level(title)
            section_stack = section_stack[: max(level - 1, 0)]
            section_stack.append(title)
            path = list(section_stack)
        else:
            level = 6
            path = [title]

        entries.append(
            PdfTocEntry(
                title=title,
                target_page=target_page,
                level=level,
                path=path,
                entry_type=current_mode,
            )
        )

    for page in pages:
        if not is_pdf_toc_page(page.page_number, page.lines or page.raw_lines):
            continue

        pending = ""
        for raw_line in page.lines or page.raw_lines:
            line = clean_heading_text(raw_line)
            if not line or line == "目录":
                continue
            if line == "表格":
                current_mode = "table"
                pending = ""
                continue
            if line == "插图":
                current_mode = "figure"
                pending = ""
                continue

            parsed_current = parse_pdf_toc_line(line, current_mode)
            if pending and parsed_current and is_probable_toc_fragment(pending):
                pending_title = normalize_pdf_toc_title(pending, current_mode)
                append_entry(pending_title, parsed_current[1])
                append_entry(parsed_current[0], parsed_current[1])
                pending = ""
                continue

            candidate = f"{pending} {line}".strip() if pending else line
            parsed = parse_pdf_toc_line(candidate, current_mode)
            if not parsed and pending:
                pending = candidate if is_probable_toc_fragment(candidate) else ""
                continue
            if not parsed:
                pending = candidate if is_probable_toc_fragment(candidate) else ""
                continue

            title, target_page = parsed
            append_entry(title, target_page)
            pending = ""

    return entries


def build_pdf_page_section_map(entries: list[PdfTocEntry], page_count: int) -> dict[int, list[str]]:
    page_map: dict[int, list[str]] = {}
    for entry in entries:
        if entry.entry_type != "section":
            continue
        page_map.setdefault(entry.target_page, list(entry.path))
    return page_map


def build_pdf_table_title_map(entries: list[PdfTocEntry]) -> dict[int, list[str]]:
    table_titles: dict[int, list[str]] = defaultdict(list)
    for entry in entries:
        if entry.entry_type == "table":
            table_titles[entry.target_page].append(entry.title)
    return dict(table_titles)


def extract_numeric_heading_key(text: str) -> str | None:
    stripped = clean_heading_text(text)
    numbered = re.match(r"^(\d+(?:\.\d+)*)", stripped)
    if numbered:
        return numbered.group(1)
    appendix = re.match(r"^(?:附录\s*)?([A-Z])(?:\.(\d+(?:\.\d+)*)?)?", stripped, flags=re.IGNORECASE)
    if not appendix:
        return None
    suffix = appendix.group(2)
    return appendix.group(1).upper() if not suffix else f"{appendix.group(1).upper()}.{suffix}"


def build_pdf_section_lookup(entries: list[PdfTocEntry]) -> dict[str, tuple[list[str], int]]:
    lookup: dict[str, tuple[list[str], int]] = {}
    for entry in entries:
        if entry.entry_type != "section":
            continue
        key = extract_numeric_heading_key(entry.title)
        if key and key not in lookup:
            lookup[key] = (list(entry.path), entry.target_page)
    return lookup


def extract_pdf_table_key(text: str) -> str | None:
    normalized = clean_heading_text(text)
    normalized = re.sub(r"([-.]\d)\s*[Oo](?=[\s.])", lambda match: f"{match.group(1)}0", normalized)
    match = re.search(r"(?:表|table)?\s*(\d+[-.]\d+)", normalized, flags=re.IGNORECASE)
    if not match:
        return None
    return match.group(1).replace(".", "-")


def canonicalize_pdf_table_title(
    title: str,
    page_number: int,
    page_table_titles: dict[int, list[str]],
) -> str:
    title = clean_heading_text(title)
    table_key = extract_pdf_table_key(title)
    candidate_pages = (page_number, page_number - 1, page_number + 1)

    if table_key:
        for candidate_page in candidate_pages:
            for candidate in page_table_titles.get(candidate_page, []):
                if extract_pdf_table_key(candidate) == table_key:
                    return candidate

    normalized = re.sub(r"^表?\s*(\d+[-.]\d+)\s*", r"表\1. ", title)
    return normalized.strip()


def derive_pdf_table_section(
    section_path: list[str],
    table_title: str,
    section_lookup: dict[str, tuple[list[str], int]] | None = None,
    page_number: int | None = None,
) -> list[str]:
    table_key = extract_pdf_table_key(table_title)
    if table_key and section_lookup:
        candidate = section_lookup.get(table_key.replace("-", "."))
        if candidate:
            candidate_path, candidate_page = candidate
            if page_number is None or candidate_page in {page_number, page_number - 1}:
                return list(candidate_path) + [table_title]

    parent = list(section_path)
    if parent and parent[-1].startswith("表"):
        parent = parent[:-1]
    return parent + [table_title]


def split_pdf_table_rows(rows: list[str]) -> tuple[str | None, list[str], list[str]]:
    cleaned_rows = [clean_heading_text(row) for row in rows if clean_heading_text(row)]
    if not cleaned_rows:
        return None, [], []

    title = cleaned_rows[0] if is_pdf_table_title(cleaned_rows[0]) else None
    remaining = cleaned_rows[1:] if title else cleaned_rows
    header_rows: list[str] = []
    body_start = 0

    for index, row in enumerate(remaining[:3]):
        if looks_like_table_header_row(row):
            header_rows.append(row)
            body_start = index + 1
            continue
        break

    body_rows = remaining[body_start:]
    if not body_rows and remaining:
        body_rows = remaining[len(header_rows) :] or remaining[-1:]
    return title, header_rows, body_rows


def flush_pdf_table_rows(
    rows: list[str],
    *,
    section_path: list[str],
    page_number: int,
) -> list[ExtractedBlock]:
    blocks: list[ExtractedBlock] = []
    title, header_rows, body_rows = split_pdf_table_rows(rows)
    prefix_rows = ([title] if title else []) + header_rows

    if not body_rows:
        text = "\n".join(prefix_rows).strip()
        if text:
            blocks.append(
                ExtractedBlock(
                    text=text,
                    block_type="table",
                    section_path=list(section_path),
                    page_number=page_number,
                    page_label=f"第 {page_number} 页",
                    location_label=f"第 {page_number} 页",
                )
            )
        return blocks

    for start in range(0, len(body_rows), 5):
        chunk_rows = prefix_rows + body_rows[start : start + 5]
        text = "\n".join(chunk_rows).strip()
        if text:
            blocks.append(
                ExtractedBlock(
                    text=text,
                    block_type="table",
                    section_path=list(section_path),
                    page_number=page_number,
                    page_label=f"第 {page_number} 页",
                    location_label=f"第 {page_number} 页",
                )
            )
    return blocks


def should_continue_pdf_table(
    carry_table_section: list[str] | None,
    base_section: list[str],
    lines: list[str],
) -> bool:
    if not carry_table_section or not carry_table_section[-1].startswith("表"):
        return False

    carry_parent = carry_table_section[:-1]
    preview = [clean_heading_text(line) for line in lines[:5] if clean_heading_text(line)]
    if not preview:
        return False
    if any(extract_inline_table_title(line)[1] for line in preview):
        return False

    table_like = sum(1 for line in preview if looks_like_table_row(line))
    if table_like < max(1, len(preview) // 2):
        return False

    if base_section and carry_parent != base_section:
        headings = [detect_pdf_heading(line) for line in preview]
        headings = [heading for heading in headings if heading and not looks_like_table_row(heading)]
        return not headings

    return True


def parse_pdf_content_page(
    page_number: int,
    lines: list[str],
    base_section: list[str],
    carry_table_section: list[str] | None,
    page_table_titles: dict[int, list[str]],
    section_lookup: dict[str, tuple[list[str], int]],
) -> tuple[list[ExtractedBlock], list[str], list[str] | None]:
    blocks: list[ExtractedBlock] = []
    text_paragraphs: list[tuple[str, dict]] = []
    table_rows: list[str] = []
    active_section = list(base_section)
    active_table_section: list[str] | None = (
        list(carry_table_section)
        if should_continue_pdf_table(carry_table_section, base_section, lines)
        else None
    )
    page_ended_in_table = False

    def flush_text():
        nonlocal text_paragraphs
        if text_paragraphs:
            blocks.extend(build_paragraph_blocks(text_paragraphs, max_chars=1200, max_items=8))
            text_paragraphs = []

    def flush_table():
        nonlocal table_rows, active_table_section
        if active_table_section and table_rows:
            blocks.extend(
                flush_pdf_table_rows(
                    table_rows,
                    section_path=active_table_section,
                    page_number=page_number,
                )
            )
        table_rows = []
        active_table_section = None

    for raw_line in lines:
        line = clean_heading_text(raw_line)
        repeated_page_heading = detect_pdf_heading(line)
        if (
            active_table_section
            and repeated_page_heading
            and heading_level(repeated_page_heading) == 1
            and not is_pdf_table_title(line)
        ):
            continue

        if active_table_section and looks_like_table_row(line):
            table_rows.append(line)
            continue

        inline_prefix, inline_table_title = extract_inline_table_title(line)
        if inline_table_title:
            flush_text()
            flush_table()
            if inline_prefix and detect_pdf_heading(inline_prefix):
                active_section = update_section_path(active_section, inline_prefix)
            canonical_title = canonicalize_pdf_table_title(
                inline_table_title,
                page_number,
                page_table_titles,
            )
            active_table_section = derive_pdf_table_section(
                active_section,
                canonical_title,
                section_lookup,
                page_number,
            )
            table_rows.append(canonical_title)
            continue

        heading = detect_pdf_heading(line)
        if heading and active_table_section:
            flush_table()

        if heading and not looks_like_table_row(line):
            flush_text()
            active_section = update_section_path(active_section, heading)
            continue

        if active_table_section:
            table_rows.append(line)
            continue

        text_paragraphs.append(
            (
                line,
                {
                    "block_type": "text",
                    "section_path": list(active_section),
                    "page_number": page_number,
                    "page_label": f"第 {page_number} 页",
                    "location_label": f"第 {page_number} 页",
                },
            )
        )

    next_table_section = list(active_table_section) if active_table_section and table_rows else None
    flush_text()
    flush_table()
    return blocks, active_section, next_table_section


def markdown_heading(line: str) -> tuple[int, str] | None:
    match = re.match(r"^(#{1,6})\s+(.+)$", line)
    if match:
        return len(match.group(1)), clean_heading_text(match.group(2))

    heading = detect_pdf_heading(line)
    if heading and not is_pdf_table_title(heading):
        return heading_level(heading), heading
    return None


def update_markdown_section_path(current: list[str], level: int, title: str) -> list[str]:
    cleaned = clean_heading_text(title)
    key = extract_numeric_heading_key(cleaned)
    if key:
        key_parts = key.split(".")
        base: list[str] = []
        for existing in current:
            existing_key = extract_numeric_heading_key(existing)
            if not existing_key:
                continue
            existing_parts = existing_key.split(".")
            if len(existing_parts) < len(key_parts) and key_parts[: len(existing_parts)] == existing_parts:
                base.append(existing)
        base.append(cleaned)
        return base

    base = current[: max(level - 1, 0)]
    base.append(cleaned)
    return base


def normalize_pdf_markdown_paragraph(lines: list[str]) -> str:
    cleaned = [clean_pdf_markdown_line(line) for line in lines if clean_pdf_markdown_line(line)]
    if not cleaned:
        return ""
    if all(line.startswith(("-", "*")) for line in cleaned):
        return "\n".join(cleaned)
    pieces: list[str] = []
    for line in cleaned:
        if line.startswith(("-", "*")):
            pieces.append("\n" + line)
        elif pieces and pieces[-1].endswith(("-", "/", "–")):
            pieces[-1] = pieces[-1] + line
        elif pieces:
            pieces.append(" " + line)
        else:
            pieces.append(line)
    text = "".join(pieces)
    text = re.sub(r"\s+", " ", text)
    text = re.sub(r"\s+\n", "\n", text)
    return text.strip()


def pdf_markdown_quality_score(blocks: list[ExtractedBlock]) -> float:
    text = "\n".join(block.text for block in blocks)
    if not text.strip():
        return 0.0
    tokens = re.findall(r"[A-Za-z][A-Za-z0-9’'_-]{3,}", text)
    if not tokens:
        return text_readability_score(text)

    glued = 0
    for token in tokens:
        if re.search(r"[a-z][A-Z][a-z]", token):
            glued += 1
        elif len(token) >= 24 and token.isalpha():
            glued += 1
    glued_ratio = glued / len(tokens)
    return text_readability_score(text) - glued_ratio * 0.9


def parse_pdf_markdown_pages(page_chunks: list[dict]) -> tuple[list[ExtractedBlock], str]:
    blocks: list[ExtractedBlock] = []
    all_texts: list[str] = []
    current_section: list[str] = []

    for chunk in page_chunks:
        metadata = chunk.get("metadata") if isinstance(chunk, dict) else {}
        page_number = int((metadata or {}).get("page") or len(all_texts) + 1)
        page_label = f"第 {page_number} 页"
        raw_text = str(chunk.get("text") or "")
        all_texts.append(raw_text)
        raw_lines = raw_text.splitlines()
        if is_probable_markdown_toc_page(page_number, raw_lines):
            continue

        paragraph_lines: list[str] = []
        table_title: str | None = None
        table_lines: list[str] = []

        def flush_paragraph() -> None:
            nonlocal paragraph_lines
            text = normalize_pdf_markdown_paragraph(paragraph_lines)
            if text:
                blocks.extend(
                    build_paragraph_blocks(
                        [
                            (
                                text,
                                {
                                    "block_type": "text",
                                    "section_path": list(current_section),
                                    "page_number": page_number,
                                    "page_label": page_label,
                                    "location_label": page_label,
                                },
                            )
                        ],
                        max_chars=1400,
                        max_items=1,
                    )
                )
            paragraph_lines = []

        def flush_table() -> None:
            nonlocal table_title, table_lines
            if not table_lines:
                table_title = None
                return
            title = table_title or "Table"
            text = "\n".join([title] + table_lines).strip()
            table_section = list(current_section)
            if title and (not table_section or table_section[-1] != title):
                table_section = table_section + [title]
            blocks.append(
                ExtractedBlock(
                    text=text,
                    block_type="table",
                    section_path=table_section,
                    page_number=page_number,
                    page_label=page_label,
                    location_label=page_label,
                )
            )
            table_title = None
            table_lines = []

        for raw_line in raw_lines:
            line = clean_pdf_markdown_line(raw_line)

            if not line:
                if table_lines:
                    flush_table()
                continue

            if line.startswith("|") and line.endswith("|"):
                flush_paragraph()
                table_lines.append(line)
                continue

            if table_lines:
                flush_table()

            heading = markdown_heading(line)
            if heading:
                flush_paragraph()
                level, title = heading
                if current_section and title.lower().startswith("chapter "):
                    continue
                current_section = update_markdown_section_path(current_section, level, title)
                table_title = None
                continue

            if is_pdf_markdown_noise_line(line):
                continue

            if is_pdf_table_title(line):
                flush_paragraph()
                table_title = line
                continue

            if re.match(r"^(Figure|图)\s+\d+[-.]\d+", line, flags=re.IGNORECASE):
                flush_paragraph()
                blocks.append(
                    ExtractedBlock(
                        text=line,
                        block_type="text",
                        section_path=list(current_section),
                        page_number=page_number,
                        page_label=page_label,
                        location_label=page_label,
                    )
                )
                continue

            paragraph_lines.append(line)

        flush_paragraph()
        flush_table()

    return postprocess_pdf_blocks(blocks), "\n".join(all_texts)


def extract_pdf_document_with_pymupdf4llm(path: Path) -> ExtractedDocument | None:
    try:
        import pymupdf4llm
    except Exception:
        return None

    try:
        page_chunks = pymupdf4llm.to_markdown(
            str(path),
            page_chunks=True,
            ignore_images=True,
            show_progress=False,
            table_strategy="lines_strict",
        )
    except Exception:
        return None

    if not isinstance(page_chunks, list):
        return None

    blocks, text = parse_pdf_markdown_pages(page_chunks)
    if pdf_markdown_quality_score(blocks) < 0.35:
        return None

    return ExtractedDocument(
        text=normalize_text(text),
        blocks=blocks,
        source_format="pdf+pymupdf4llm",
        page_count=len(page_chunks),
    )


def extract_pdf_document(path: Path) -> ExtractedDocument:
    markdown_document = extract_pdf_document_with_pymupdf4llm(path)
    if markdown_document and markdown_document.blocks:
        return markdown_document

    reader = PdfReader(str(path))
    blocks: list[ExtractedBlock] = []
    all_texts: list[str] = []
    used_ocr = False
    pages: list[PdfPageData] = []

    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text(extraction_mode="layout") or page.extract_text() or ""
        text = normalize_text(text)
        text, lines, page_used_ocr = select_pdf_page_text(path, index, text)
        used_ocr = used_ocr or page_used_ocr
        if not text:
            continue

        all_texts.append(text)
        pages.append(
            PdfPageData(
                page_number=index,
                page_label=f"第 {index} 页",
                text=text,
                raw_lines=lines,
                used_ocr=page_used_ocr,
            )
        )

    header_noise, footer_noise = collect_pdf_margin_noise(pages)
    for page in pages:
        page.lines = strip_pdf_page_noise(page, header_noise=header_noise, footer_noise=footer_noise)

    toc_entries = build_pdf_toc_entries(pages)
    page_section_map = build_pdf_page_section_map(toc_entries, len(reader.pages))
    page_table_titles = build_pdf_table_title_map(toc_entries)
    section_lookup = build_pdf_section_lookup(toc_entries)
    current_section: list[str] = []
    carry_table_section: list[str] | None = None

    for page in pages:
        if is_pdf_toc_page(page.page_number, page.lines or page.raw_lines):
            blocks.extend(parse_pdf_toc_page(page.page_number, page.lines or page.raw_lines))
            continue

        fallback_section = current_section[:-1] if current_section and current_section[-1].startswith("表") else current_section
        base_section = page_section_map.get(page.page_number) or fallback_section
        page_blocks, current_section, carry_table_section = parse_pdf_content_page(
            page.page_number,
            page.lines or page.raw_lines,
            base_section,
            carry_table_section,
            page_table_titles,
            section_lookup,
        )
        blocks.extend(page_blocks)
    blocks = postprocess_pdf_blocks(blocks)
    return ExtractedDocument(
        text="\n".join(all_texts),
        blocks=blocks,
        source_format="pdf+ocr" if used_ocr else "pdf",
        page_count=len(reader.pages),
    )


def extract_docx_document(path: Path) -> ExtractedDocument:
    document = DocxDocument(str(path))
    section_path: list[str] = []
    paragraphs: list[tuple[str, dict]] = []
    all_texts: list[str] = []

    for index, paragraph in enumerate(document.paragraphs, start=1):
        text = normalize_text(paragraph.text)
        if not text:
            continue

        style_name = (paragraph.style.name or "").lower()
        heading_match = re.search(r"heading\s*([1-6])|标题\s*([1-6])", style_name)
        if heading_match:
            level = int(heading_match.group(1) or heading_match.group(2))
            section_path = section_path[: level - 1]
            section_path.append(text)
            continue

        all_texts.append(text)
        paragraphs.append(
            (
                text,
                {
                    "section_path": list(section_path),
                    "paragraph_start": index,
                    "paragraph_end": index,
                    "location_label": f"段落 {index}",
                },
            )
        )

    blocks = build_paragraph_blocks(paragraphs)
    return ExtractedDocument(
        text="\n".join(all_texts),
        blocks=blocks,
        source_format="docx",
    )


def extract_markdown_document(path: Path, text: str, encoding: str | None) -> ExtractedDocument:
    text = strip_sphinx_boilerplate(text)
    lines = text.splitlines()
    blocks = build_line_blocks(lines, filename=path.name, source_uri=path.name)
    return ExtractedDocument(
        text=text,
        blocks=blocks,
        source_format="markdown",
        encoding=encoding,
    )


def extract_code_document(path: Path, text: str, encoding: str | None) -> ExtractedDocument:
    lines = text.splitlines()
    blocks = build_line_blocks(lines, filename=path.name, source_uri=path.name, block_type="code")
    return ExtractedDocument(
        text=text,
        blocks=blocks,
        source_format="code",
        encoding=encoding,
    )


def extract_plain_text_document(path: Path, text: str, encoding: str | None) -> ExtractedDocument:
    text = strip_sphinx_boilerplate(text)
    lines = text.splitlines()
    blocks = build_line_blocks(lines, filename=path.name, source_uri=path.name)
    return ExtractedDocument(
        text=text,
        blocks=blocks,
        source_format="text",
        encoding=encoding,
    )


def extract_document_from_path(path: Path) -> ExtractedDocument:
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return extract_pdf_document(path)

    if suffix in WORD_SUFFIXES:
        return extract_docx_document(path)

    if suffix in TEXT_SUFFIXES | MARKDOWN_SUFFIXES | CODE_SUFFIXES:
        raw = path.read_bytes()
        text, encoding = decode_text_bytes(raw)
        if suffix in MARKDOWN_SUFFIXES:
            return extract_markdown_document(path, text, encoding)
        if suffix in CODE_SUFFIXES:
            return extract_code_document(path, text, encoding)
        return extract_plain_text_document(path, text, encoding)

    raise ValueError("仅支持 .txt / .md / .pdf / .docx 及常见代码文件。")


def extract_text_from_path(path: Path) -> str:
    return extract_document_from_path(path).text


class MainContentHTMLParser(HTMLParser):
    def __init__(self):
        super().__init__()
        self.capture_depth = 0
        self.capture_stack: list[bool] = []
        self.skip_depth = 0
        self.pre_depth = 0
        self.text_parts: list[str] = []
        self.title_text: list[str] = []
        self._in_title = False

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]):
        attrs_dict = dict(attrs)
        class_attr = attrs_dict.get("class") or ""
        class_tokens = {token.strip().lower() for token in class_attr.split() if token.strip()}
        role_attr = attrs_dict.get("role") or ""
        id_attr = attrs_dict.get("id") or ""

        if tag in {"script", "style", "noscript"}:
            self.skip_depth += 1
            return

        if tag == "title":
            self._in_title = True

        should_capture = (
            tag in {"main", "article"}
            or role_attr == "main"
            or id_attr in {"main-content", "content"}
            or bool({"document", "main", "content", "app-content"} & class_tokens)
        )
        self.capture_stack.append(should_capture)
        if should_capture:
            self.capture_depth += 1

        if self.capture_depth and tag == "pre":
            self.pre_depth += 1
            self.text_parts.append("\n```\n")
            return
        if self.capture_depth and tag == "br":
            self.text_parts.append("\n")
            return
        if self.capture_depth and tag in {"p", "div", "section", "article", "li", "tr"}:
            self.text_parts.append("\n")
        if self.capture_depth and tag in {"h1", "h2", "h3", "h4", "h5", "h6"}:
            level = int(tag[1])
            self.text_parts.append(f"\n{'#' * level} ")

    def handle_endtag(self, tag: str):
        if tag in {"script", "style", "noscript"} and self.skip_depth:
            self.skip_depth -= 1
            return

        if tag == "title":
            self._in_title = False

        if tag == "pre" and self.pre_depth:
            self.pre_depth -= 1
            self.text_parts.append("\n```\n")

        if self.capture_stack:
            should_capture = self.capture_stack.pop()
        else:
            should_capture = False

        if should_capture and self.capture_depth:
            self.capture_depth -= 1

        if self.capture_depth and tag in {"p", "div", "section", "article", "li", "tr"}:
            self.text_parts.append("\n")
        if self.capture_depth and tag in {"h1", "h2", "h3", "h4", "h5", "h6"}:
            self.text_parts.append("\n")

    def handle_data(self, data: str):
        if self._in_title:
            self.title_text.append(data)

        if self.skip_depth or not self.capture_depth:
            return

        if self.pre_depth:
            self.text_parts.append(data)
            return

        cleaned = re.sub(r"\s+", " ", data).strip()
        if cleaned:
            self.text_parts.append(cleaned)

    @property
    def extracted_title(self) -> str:
        return re.sub(r"\s+", " ", " ".join(self.title_text)).strip()

    @property
    def extracted_text(self) -> str:
        return normalize_text("".join(self.text_parts))


def extract_web_document(url: str) -> tuple[str, ExtractedDocument]:
    request = Request(
        url,
        headers={
            "User-Agent": "Mozilla/5.0 (compatible; DeepSeekRAG/0.1; +https://api.deepseek.com)",
        },
    )
    with urlopen(request, timeout=30) as response:
        raw = response.read()
        html, encoding = decode_text_bytes(raw)

    parser = MainContentHTMLParser()
    parser.feed(html)
    text = strip_sphinx_boilerplate(parser.extracted_text)
    if not text:
        raise ValueError("网页正文提取失败，请尝试换一个更具体的文档页面。")

    title = parser.extracted_title or urlparse(url).path.rsplit("/", 1)[-1] or url
    title = re.sub(r"\s+", " ", title).strip()

    blocks = build_line_blocks(
        text.splitlines(),
        filename=title,
        source_uri=url,
        max_lines=60,
        max_chars=1200,
    )
    for block in blocks:
        if block.section_path:
            block.source_uri = f"{url}#{slugify(block.section_path[-1])}"
            block.location_label = block.section_path[-1]
        elif not block.location_label:
            block.location_label = url

    return title, ExtractedDocument(
        text=text,
        blocks=blocks,
        source_format="web",
        encoding=encoding,
    )


def fetch_web_document(url: str) -> tuple[str, str]:
    title, document = extract_web_document(url)
    return title, document.text
